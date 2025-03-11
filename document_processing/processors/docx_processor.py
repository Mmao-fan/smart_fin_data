# -*- coding: utf-8 -*-
# processors/docx_processor.py
import logging
from pathlib import Path
from .base_processor import BaseProcessor
from .exceptions import DocumentProcessingError
from typing import Dict, Any

# 尝试导入 python-docx，如果不存在则提供一个替代方案
try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx 模块未安装，Word处理功能将受限")


class DocxProcessor(BaseProcessor):
    @classmethod
    def extract_text(cls, file_path: str) -> Dict[str, Any]:
        """从Word文档中提取文本"""
        try:
            if not HAS_DOCX:
                raise DocumentProcessingError("python-docx 模块未安装，无法处理Word文件。请安装 python-docx: pip install python-docx")
                
            doc = Document(file_path)
            
            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格文本
            tables = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if any(row_text):  # 只添加非空行
                        table_text.append(' | '.join(row_text))
                if table_text:
                    tables.append('\n'.join(table_text))
            
            # 合并所有文本
            text_chunks = paragraphs + tables
            full_text = '\n\n'.join(text_chunks)
            
            # 收集元数据
            metadata = {
                'sections': len(doc.sections),
                'paragraphs': len(paragraphs),
                'tables': len(tables),
                'has_headers': any(section.header for section in doc.sections),
                'has_footers': any(section.footer for section in doc.sections)
            }
            
            return {
                'text': full_text,
                'text_chunks': text_chunks,
                'total_pages': len(text_chunks),  # 使用块数作为页数的估计
                'metadata': metadata
            }
            
        except Exception as e:
            cls.safe_logging(file_path, e)
            if isinstance(e, DocumentProcessingError):
                raise
            raise DocumentProcessingError(f"处理Word文档失败: {str(e)}")

    @classmethod
    def _parse_table(cls, table) -> str:
        """优化表格空值处理"""
        if not HAS_DOCX:
            return ""
            
        markdown = []
        for i, row in enumerate(table.rows):
            cells = [
                (cell.text.strip().replace('\n', ' ') or "[空]")  # 处理空单元格
                for cell in row.cells
            ]
            markdown.append("| " + " | ".join(cells) + " |")
            if i == 0:  # 添加分隔行
                markdown.append("| " + " | ".join(["---"] * len(row.cells)) + " |")
        return "\n".join(markdown)