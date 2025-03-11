# -*- coding: utf-8 -*-
import os
import logging
import pandas as pd
import json
from typing import Dict, List, Any, Optional, Union, Generator
import traceback
import re
import sys
import importlib
from datetime import datetime
from io import StringIO
import numpy as np
from collections import defaultdict
import subprocess
from pathlib import Path
import time
from document_processing import DocumentProcessor
from text_chunking import ChunkManager
from information_extraction import InformationProcessor, EnhancedAdaptiveSystem

# 添加项目根目录到系统路径
project_root = os.path.dirname(os.path.abspath(__file__))
if project_root not in sys.path:
    sys.path.append(project_root)

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('processing.log', mode='w'),
        logging.StreamHandler()
    ]
)

# 设置特定模块的日志级别
logging.getLogger('PIL').setLevel(logging.WARNING)
logging.getLogger('matplotlib').setLevel(logging.WARNING)
logging.getLogger('transformers').setLevel(logging.WARNING)

logger = logging.getLogger(__name__)

# 导入自定义模块
try:
    from document_processing.document_processor import DocumentProcessor
    from text_chunking.chunk_manager import ChunkManager
    from information_extraction.information_extractor import InformationProcessor
    from information_extraction.enhanced_adaptive_system import EnhancedAdaptiveSystem
    from information_extraction.anomaly_detector import AnomalyDetector
    from scenario_adaptation.customer_service_generator import CustomerServiceGenerator
    from scenario_adaptation.fraud_encoder import FraudEncoder
    from scenario_adaptation.compliance_mapper import ComplianceMapper
    from information_extraction.relation_extractor import RelationExtractor
    logger.info("成功导入所有自定义模块")
except ImportError as e:
    logger.error(f"导入自定义模块失败: {str(e)}")
    raise  # 如果模块导入失败，直接抛出异常，因为这些模块是必需的

def is_module_installed(module_name: str) -> bool:
    """检查模块是否已安装"""
    try:
        importlib.import_module(module_name)
        return True
    except ImportError:
        return False

def setup_logging():
    """设置日志配置"""
    # 创建日志目录
    log_dir = os.path.join(project_root, 'logs')
    os.makedirs(log_dir, exist_ok=True)
    
    # 配置日志格式
    log_format = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    logging.basicConfig(
        level=logging.INFO,
        format=log_format,
        handlers=[
            logging.FileHandler(os.path.join(log_dir, 'processing.log'), mode='w'),
            logging.StreamHandler()
        ]
    )
    
    # 设置第三方库的日志级别
    logging.getLogger('PIL').setLevel(logging.WARNING)
    logging.getLogger('matplotlib').setLevel(logging.WARNING)
    logging.getLogger('transformers').setLevel(logging.WARNING)

# 设置pip安装源
PIP_INDEX_URL = "https://mirrors.aliyun.com/pypi/simple/"
PIP_TRUSTED_HOST = "mirrors.aliyun.com"

def install_module(module_name):
    """安装Python模块"""
    try:
        logging.info(f"正在安装 {module_name} 模块...")
        cmd = [
            sys.executable,
            "-m",
            "pip",
            "install",
            "--index-url",
            PIP_INDEX_URL,
            "--trusted-host",
            PIP_TRUSTED_HOST,
            module_name
        ]
        subprocess.check_call(cmd)
        logging.info(f"{module_name} 模块安装成功")
        return True
    except Exception as e:
        logging.error(f"{module_name} 模块安装失败: {str(e)}")
        return False

# 检查并安装必要的依赖
required_modules = {
    'pandas': 'pandas>=1.3.0',
    'python-docx': 'python-docx>=0.8.11',
    'PyPDF2': 'PyPDF2>=2.0.0',
    'scikit-learn': 'scikit-learn>=0.24.2',
    'numpy': 'numpy>=1.21.0'
}

# 尝试导入必要模块，如果失败则进入离线模式
OFFLINE_MODE = False
try:
    import pandas as pd
    import numpy as np
    from docx import Document
    import PyPDF2
    from sklearn.feature_extraction.text import TfidfVectorizer
    logger.info("成功加载核心依赖")
except ImportError as e:
    logger.warning(f"无法加载某些核心依赖，进入离线模式: {str(e)}")
    OFFLINE_MODE = True

# 尝试导入可选模块
try:
    from sentence_transformers import SentenceTransformer
    HAS_SENTENCE_TRANSFORMERS = True
    logger.info("成功加载sentence-transformers")
except ImportError:
    HAS_SENTENCE_TRANSFORMERS = False
    logger.info("sentence-transformers未安装，将使用基础文本处理方法")

try:
    import torch
    HAS_TORCH = True
    logger.info("成功加载PyTorch")
except ImportError:
    HAS_TORCH = False
    logger.info("PyTorch未安装，将使用CPU处理")

# 检查并安装 python-docx 模块
HAS_DOCX = not OFFLINE_MODE and is_module_installed('docx')
if not HAS_DOCX:
    logger.warning("python-docx 模块未安装，尝试自动安装...")
    if install_module('python-docx'):
        HAS_DOCX = True
        from docx import Document
    else:
        logger.warning("python-docx 模块安装失败，Word处理功能将受限")
else:
    logger.info("python-docx 模块已安装")
    from docx import Document

# 检查并安装 PyPDF2
HAS_PYPDF2 = not OFFLINE_MODE and is_module_installed('PyPDF2')
if not HAS_PYPDF2:
    logger.warning("PyPDF2 模块未安装，尝试自动安装...")
    if install_module('PyPDF2'):
        HAS_PYPDF2 = True
        import PyPDF2
    else:
        logger.warning("PyPDF2 模块未安装，PDF处理功能将受限")
else:
    logger.info("PyPDF2 模块已安装")
    import PyPDF2

# 不安装spacy，使用离线模式
HAS_SPACY = False
logger.info("使用离线模式进行实体识别和文本处理")

# 检查并安装必要的依赖
required_modules = {
    'pandas': 'pandas',
    'python-docx': 'docx',
    'PyPDF2': 'PyPDF2',
    'scikit-learn': 'sklearn'
}

for module_name, import_name in required_modules.items():
    if not is_module_installed(import_name):
        logger.warning(f"{module_name} 模块未安装，尝试自动安装...")
        if install_module(module_name):
            logger.info(f"{module_name} 模块安装成功")
        else:
            logger.warning(f"{module_name} 模块安装失败，部分功能可能受限")
    else:
        logger.info(f"{module_name} 模块已安装")

# 尝试导入其他模块，如果导入失败则提供替代方案
try:
    from document_processing.document_processor import DocumentProcessor

    HAS_DOCUMENT_PROCESSOR = True
except ImportError:
    HAS_DOCUMENT_PROCESSOR = False
    logger.warning("DocumentProcessor 模块导入失败，只能处理CSV文件、Word文件和PDF文件")

try:
    from text_chunking.chunk_manager import ChunkManager
    from information_extraction.anomaly_detector import AnomalyDetector
    from information_extraction.information_extractor import InformationProcessor

    HAS_TEXT_PROCESSING = True
except ImportError:
    HAS_TEXT_PROCESSING = False
    logger.warning("文本处理模块导入失败，部分功能将受限")

try:
    from scenario_adaptation.customer_service_generator import CustomerServiceGenerator
    from scenario_adaptation.fraud_encoder import FraudEncoder
    from scenario_adaptation.compliance_mapper import ComplianceMapper
    from config import CUSTOMER_SERVICE_CONFIG

    HAS_SCENARIO_ADAPTATION = True
except ImportError:
    HAS_SCENARIO_ADAPTATION = False
    logger.warning("场景适配模块导入失败，部分功能将受限")


def detect_file_encoding(file_path: str) -> str:
    """检测文件编码"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'iso-8859-1', 'ascii']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                f.read(100)  # 尝试读取前100个字符
                return encoding
        except UnicodeDecodeError:
            continue
    
    # 如果所有编码都失败，默认返回utf-8
    return 'utf-8'


def chunk_large_file(file_path: str, chunk_size: int = 1024 * 1024) -> Generator[str, None, None]:
    """分块读取大文件"""
    encoding = detect_file_encoding(file_path)
    with open(file_path, 'r', encoding=encoding) as file:
        while True:
            chunk = file.read(chunk_size)
            if not chunk:
                break
            yield chunk


def extract_advanced_keywords(text: str) -> Dict[str, List[str]]:
    """增强的关键词提取"""
    keywords = defaultdict(list)

    # 使用更复杂的模式匹配
    patterns = {
        'banks': [
            r"([A-Za-z\s]+(?:Bank|Financial|Credit Union))",
            r"([\u4e00-\u9fa5]+(?:银行|信用社|金融))",
        ],
        'companies': [
            r"([A-Z][a-z]+(?:\s[A-Z][a-z]+)*\s(?:Inc|Corp|Ltd|LLC|Company|Group))",
            r"([\u4e00-\u9fa5]+(?:公司|集团|企业|有限责任|股份))",
        ],
        'dates': [
            r"(\d{4}(?:/\d{1,2}){2})",
            r"(\d{4}年\d{1,2}月\d{1,2}日)",
            r"(\d{1,2}/\d{1,2}/\d{4})",
            r"(\d{4}-\d{2}-\d{2})",
        ],
        'amounts': [
            r"(\d+(?:\.\d+)?)\s*(?:亿|万|元|美元|USD|CNY|RMB|€|₤|¥)",
            r"(?:USD|CNY|RMB|€|₤|¥)\s*(\d+(?:\.\d+)?)",
        ],
        'locations': [
            r"([\u4e00-\u9fa5]{2,}(?:省|市|区|县|镇))",
            r"([A-Z][a-z]+(?:\s[A-Z][a-z]+)*(?:\s+City)?)",
        ],
        'emails': [
            r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
        ],
        'phones': [
            r"(\+?\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4})",
            r"(\d{3,4}[-\s]?\d{3,4}[-\s]?\d{4})",
        ],
    }

    for category, pattern_list in patterns.items():
        for pattern in pattern_list:
            matches = re.finditer(pattern, text)
            for match in matches:
                value = match.group(1).strip()
                if value and value not in keywords[category]:
                    keywords[category].append(value)

    return dict(keywords)


def extract_document_structure(content: str) -> Dict[str, Any]:
    """提取文档结构"""
    structure = {
        'paragraphs': [],
        'sections': [],
        'lists': [],
        'tables': [],
    }

    # 分析段落
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', content) if p.strip()]
    structure['paragraphs'] = paragraphs

    # 识别章节
    section_pattern = r'^(?:第[一二三四五六七八九十]+[章节]|[IVX]+\.|[\d]+\.)\s*(.+)$'
    for para in paragraphs:
        if re.match(section_pattern, para, re.MULTILINE):
            structure['sections'].append(para)

    # 识别列表
    list_pattern = r'(?:^[\d]+\.|^[-•*]\s+)(.+)$'
    current_list = []
    for para in paragraphs:
        if re.match(list_pattern, para, re.MULTILINE):
            current_list.append(para)
        elif current_list:
            if len(current_list) > 1:
                structure['lists'].append(current_list.copy())
            current_list = []

    # 识别表格（简单表格）
    table_pattern = r'[|｜].+[|｜]'
    current_table = []
    for para in paragraphs:
        if re.match(table_pattern, para):
            current_table.append(para)
        elif current_table:
            if len(current_table) > 1:
                structure['tables'].append(current_table.copy())
            current_table = []

    return structure


def process_csv_file(file_path: str) -> Dict[str, Any]:
    """增强的CSV文件处理"""
    try:
        logger.info(f"处理CSV文件: {file_path}")

        # 检测文件编码
        encoding = detect_file_encoding(file_path)
        logger.info(f"检测到文件编码: {encoding}")

        # 尝试不同的分隔符
        separators = [',', ';', '\t', '|']
        df = None
        used_sep = None

        for sep in separators:
            try:
                df = pd.read_csv(file_path, encoding=encoding, sep=sep)
                used_sep = sep
                break
            except:
                continue

        if df is None:
            raise ValueError("无法识别CSV文件格式")

        # 数据清洗和预处理
        df = df.replace(['', 'null', 'NULL', 'NaN', 'nan'], np.nan)

        # 识别列类型
        column_types = {}
        for col in df.columns:
            if df[col].dtype == 'object':
                # 尝试转换为日期，支持多种日期格式
                date_formats = [
                    '%Y-%m-%d',
                    '%Y/%m/%d',
                    '%d/%m/%Y',
                    '%m/%d/%Y',
                    '%Y%m%d',
                    '%Y-%m-%d %H:%M:%S',
                    '%Y/%m/%d %H:%M:%S'
                ]
                is_date = False
                for date_format in date_formats:
                    try:
                        pd.to_datetime(df[col], format=date_format, errors='raise')
                        column_types[col] = 'date'
                        is_date = True
                        break
                    except:
                        continue

                if not is_date:
                    # 检查是否是数值（带有货币符号等）
                    try:
                        if df[col].str.contains(r'[\d]+').all():
                            column_types[col] = 'numeric_string'
                        else:
                            column_types[col] = 'text'
                    except:
                        column_types[col] = 'text'
            else:
                column_types[col] = str(df[col].dtype)

        # 构建结构化数据
        structured_data = {
            'type': 'tabular_data',
            'data': df.to_dict(orient='records'),
            'columns': df.columns.tolist(),
            'column_types': column_types,
            'row_count': len(df),
            'metadata': {
                'file_path': str(file_path),  # 转换为字符串
                'file_type': '.csv',
                'encoding': encoding,
                'separator': used_sep,
                'file_size': os.path.getsize(file_path),
                'processed_time': datetime.now().isoformat()
            },
            'statistics': {
                'missing_values': df.isnull().sum().to_dict(),
                'unique_values': {col: int(df[col].nunique()) for col in df.columns},  # 转换为整数
                'numeric_columns': df.select_dtypes(include=[np.number]).columns.tolist()
            }
        }

        return structured_data

    except Exception as e:
        logger.error(f"处理CSV文件失败: {str(e)}", exc_info=True)
        return None


def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """增强的Word文档文本提取"""
    if not HAS_DOCX:
        return None

    try:
        doc = Document(file_path)
        document_data = {
            'paragraphs': [],
            'tables': [],
            'sections': [],
            'headers': [],
            'footers': [],
            'images': [],
            'styles': set()
        }

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                para_data = {
                    'text': para.text.strip(),
                    'style': para.style.name,
                    'alignment': str(para.alignment) if para.alignment else 'LEFT'
                }
                # 安全地获取段落格式属性
                try:
                    if hasattr(para.paragraph_format, 'first_line_indent'):
                        para_data['first_line_indent'] = para.paragraph_format.first_line_indent
                    if hasattr(para.paragraph_format, 'left_indent'):
                        para_data['left_indent'] = para.paragraph_format.left_indent
                    if hasattr(para.paragraph_format, 'right_indent'):
                        para_data['right_indent'] = para.paragraph_format.right_indent
                except:
                    pass

                document_data['paragraphs'].append(para_data)
                document_data['styles'].add(para.style.name)

        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_data = {
                        'text': cell.text.strip()
                    }
                    # 安全地获取单元格跨度
                    try:
                        if hasattr(cell._tc, 'tcPr') and cell._tc.tcPr:
                            if hasattr(cell._tc.tcPr, 'gridSpan'):
                                cell_data['spans'] = len(cell._tc.tcPr.gridSpan)
                    except:
                        cell_data['spans'] = 1
                    row_data.append(cell_data)
                table_data.append(row_data)
            document_data['tables'].append(table_data)

        # 提取节
        for section in doc.sections:
            section_data = {
                'start_type': str(section.start_type) if hasattr(section, 'start_type') else None,
                'orientation': str(section.orientation) if hasattr(section, 'orientation') else None,
                'page_height': float(section.page_height.cm) if hasattr(section, 'page_height') else None,
                'page_width': float(section.page_width.cm) if hasattr(section, 'page_width') else None
            }

            # 安全地提取页眉页脚
            try:
                if section.header and section.header.paragraphs:
                    section_data['header'] = "\n".join(
                        p.text.strip() for p in section.header.paragraphs if p.text.strip())
                else:
                    section_data['header'] = None
            except:
                section_data['header'] = None

            try:
                if section.footer and section.footer.paragraphs:
                    section_data['footer'] = "\n".join(
                        p.text.strip() for p in section.footer.paragraphs if p.text.strip())
                else:
                    section_data['footer'] = None
            except:
                section_data['footer'] = None

            document_data['sections'].append(section_data)

        document_data['styles'] = list(document_data['styles'])
        return document_data

    except Exception as e:
        logger.error(f"提取Word文档文本失败: {str(e)}")
        return None


def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """增强的PDF文档文本提取"""
    if not HAS_PYPDF2:
        logger.error("PyPDF2 模块未安装")
        return None

    try:
        document_data = {
            'pages': [],
            'metadata': {},
            'images': [],
            'forms': [],
            'links': []
        }

        logger.info(f"开始处理PDF文件: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"PDF文件不存在: {file_path}")
            return None
            
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            logger.error(f"PDF文件为空: {file_path}")
            return None
            
        logger.info(f"PDF文件大小: {file_size} 字节")

        with open(file_path, 'rb') as file:
            try:
                reader = PyPDF2.PdfReader(file)
                logger.info(f"成功创建PDF阅读器")
            except Exception as e:
                logger.error(f"创建PDF阅读器失败: {str(e)}")
                return None
            
            if not reader.pages:
                logger.warning(f"PDF文件 {file_path} 没有页面")
                return None
                
            logger.info(f"PDF文件页数: {len(reader.pages)}")

            # 提取元数据
            try:
                document_data['metadata'] = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                    'producer': reader.metadata.get('/Producer', ''),
                    'creation_date': reader.metadata.get('/CreationDate', ''),
                    'modification_date': reader.metadata.get('/ModDate', '')
                }
                logger.info("成功提取元数据")
            except Exception as e:
                logger.warning(f"提取元数据失败: {str(e)}")

            # 提取页面内容
            for page_num in range(len(reader.pages)):
                try:
                    page = reader.pages[page_num]
                    logger.info(f"处理第 {page_num + 1} 页")
                    
                    try:
                        text = page.extract_text()
                        if text:
                            logger.info(f"第 {page_num + 1} 页成功提取文本，长度: {len(text)}")
                        else:
                            logger.warning(f"第 {page_num + 1} 页文本为空")
                    except Exception as e:
                        logger.error(f"提取第 {page_num + 1} 页文本失败: {str(e)}")
                        text = ""
                    
                    # 如果页面文本为空，尝试OCR
                    if not text and HAS_TESSERACT:
                        try:
                            logger.info(f"尝试对第 {page_num + 1} 页进行OCR处理")
                            # 将PDF页面转换为图像
                            images = convert_from_path(file_path, first_page=page_num+1, last_page=page_num+1)
                            if images:
                                text = pytesseract.image_to_string(images[0], lang='chi_sim+eng')
                                if text:
                                    logger.info(f"OCR成功提取文本，长度: {len(text)}")
                                else:
                                    logger.warning("OCR未能提取到文本")
                        except Exception as e:
                            logger.warning(f"OCR处理失败: {str(e)}")
                    
                    page_data = {
                        'number': page_num + 1,
                        'text': text or "",
                        'size': {
                            'width': float(page.mediabox.width),
                            'height': float(page.mediabox.height)
                        },
                        'rotation': page.get('/Rotate', 0)
                    }

                    # 提取表单字段
                    if '/AcroForm' in page:
                        form_fields = []
                        try:
                            for annot in page['/Annots']:
                                if isinstance(annot, PyPDF2.generic.IndirectObject):
                                    annot = annot.get_object()
                                if annot.get('/FT'):
                                    field_data = {
                                        'type': str(annot['/FT']),
                                        'name': str(annot.get('/T', '')),
                                        'value': str(annot.get('/V', ''))
                                    }
                                    form_fields.append(field_data)
                            logger.info(f"成功提取 {len(form_fields)} 个表单字段")
                        except Exception as e:
                            logger.warning(f"提取表单字段失败: {str(e)}")
                        page_data['forms'] = form_fields

                    # 提取链接
                    links = []
                    try:
                        if hasattr(page, 'annotations') and page.annotations:
                            for annot in page.annotations:
                                try:
                                    if isinstance(annot, PyPDF2.generic.IndirectObject):
                                        annot = annot.get_object()
                                    if annot and annot.get('/Subtype') == '/Link' and annot.get('/A'):
                                        link_data = {
                                            'type': 'external',
                                            'url': str(annot['/A'].get('/URI', ''))
                                        }
                                        if link_data['url']:  # 只添加有效的URL
                                            links.append(link_data)
                                except Exception as e:
                                    continue  # 跳过单个链接的错误
                            logger.info(f"成功提取 {len(links)} 个链接")
                    except Exception as e:
                        logger.debug(f"页面 {page_num + 1} 没有链接或链接提取失败")
                    page_data['links'] = links

                    document_data['pages'].append(page_data)
                except Exception as e:
                    logger.error(f"处理第 {page_num + 1} 页时出错: {str(e)}")
                    continue

            # 检查是否成功提取了任何文本
            total_text = ''.join(page.get('text', '') for page in document_data['pages'])
            if not total_text.strip():
                logger.warning(f"未能从PDF文件 {file_path} 提取到任何文本")
                return None
                
            logger.info(f"成功提取文本，总长度: {len(total_text)}")

        return document_data

    except Exception as e:
        logger.error(f"提取PDF文档文本失败: {str(e)}")
        return None


def process_docx_file(file_path: str) -> Dict[str, Any]:
    """增强的Word文档处理"""
    try:
        logger.info(f"处理Word文档: {file_path}")

        # 提取文档内容和结构
        doc_data = extract_text_from_docx(file_path)
        if not doc_data:
            return None

        # 构建完整文本
        full_text = "\n".join([p['text'] for p in doc_data['paragraphs']])

        # 提取关键信息
        keywords = extract_advanced_keywords(full_text)

        # 分析文档结构
        doc_structure = extract_document_structure(full_text)

        # 构建结构化数据
        structured_data = {
            'type': 'document',
            'content': full_text[:1000] + "..." if len(full_text) > 1000 else full_text,
            'metadata': {
                'file_path': file_path,
                'file_type': '.docx',
                'file_size': os.path.getsize(file_path),
                'processed_time': datetime.now().isoformat(),
                'document_properties': {
                    'sections': len(doc_data['sections']),
                    'tables': len(doc_data['tables']),
                    'styles': doc_data['styles']
                }
            },
            'structure': {
                'paragraphs': doc_structure['paragraphs'][:100],
                'sections': doc_structure['sections'],
                'lists': doc_structure['lists'],
                'tables': doc_structure['tables'],
                'document_elements': {
                    'tables': doc_data['tables'],
                    'headers': doc_data['headers'],
                    'footers': doc_data['footers']
                }
            },
            'keywords': keywords
        }

        return structured_data

    except Exception as e:
        logger.error(f"处理Word文档失败: {str(e)}", exc_info=True)
        return None


def process_pdf_file(file_path: str) -> Dict[str, Any]:
    """增强的PDF文档处理"""
    try:
        logger.info(f"处理PDF文档: {file_path}")

        # 提取PDF内容和结构
        pdf_data = extract_text_from_pdf(file_path)
        if not pdf_data:
            return None

        # 构建完整文本
        full_text = "\n".join([page.get('text', '') for page in pdf_data['pages']])
        if not full_text.strip():
            logger.warning(f"PDF文件 {file_path} 提取的文本为空")
            return None

        # 提取关键信息
        keywords = extract_advanced_keywords(full_text)

        # 分析文档结构
        doc_structure = extract_document_structure(full_text)

        # 构建结构化数据
        structured_data = {
            'type': 'document',
            'content': full_text,  # 存储完整文本
            'metadata': {
                'file_path': file_path,
                'file_type': '.pdf',
                'file_size': os.path.getsize(file_path),
                'processed_time': datetime.now().isoformat(),
                'pdf_metadata': pdf_data['metadata'],
                'document_properties': {
                    'pages': len(pdf_data['pages']),
                    'forms': any(page.get('forms') for page in pdf_data['pages']),
                    'links': any(page.get('links') for page in pdf_data['pages'])
                }
            },
            'structure': {
                'paragraphs': doc_structure['paragraphs'][:100],
                'sections': doc_structure['sections'],
                'lists': doc_structure['lists'],
                'tables': doc_structure['tables'],
                'pdf_elements': {
                    'pages': [{
                        'number': page['number'],
                        'size': page['size'],
                        'rotation': page['rotation'],
                        'forms': page.get('forms', []),
                        'links': page.get('links', [])
                    } for page in pdf_data['pages']]
                }
            },
            'keywords': keywords
        }

        return structured_data

    except Exception as e:
        logger.error(f"处理PDF文档失败: {str(e)}", exc_info=True)
        return None


def detect_scenario(file_path: str, content: str = None) -> str:
    """
    自动检测文件场景
    使用场景适配模块进行智能场景检测
    """
    try:
        # 如果有场景适配模块，优先使用
        if HAS_SCENARIO_ADAPTATION:
            # 读取文件内容（如果未提供）
            if content is None:
                try:
                    with open(file_path, 'r', encoding=detect_file_encoding(file_path)) as f:
                        content = f.read()
                except Exception as e:
                    logger.warning(f"读取文件内容失败: {str(e)}")
                    content = ""

            # 使用各个场景检测器
            scores = {
                'customer_service': 0,
                'fraud_detection': 0,
                'compliance': 0
            }

            # 客服场景检测
            try:
                cs_generator = CustomerServiceGenerator()
                cs_score = cs_generator.evaluate_content(content)
                scores['customer_service'] = cs_score
            except Exception as e:
                logger.debug(f"客服场景评分失败: {str(e)}")

            # 欺诈检测场景检测
            try:
                fraud_encoder = FraudEncoder()
                fraud_score = fraud_encoder.analyze_risk(content)
                scores['fraud_detection'] = fraud_score
            except Exception as e:
                logger.debug(f"欺诈检测场景评分失败: {str(e)}")

            # 合规场景检测
            try:
                compliance_mapper = ComplianceMapper()
                compliance_score = compliance_mapper.evaluate_compliance(content)
                scores['compliance'] = compliance_score
            except Exception as e:
                logger.debug(f"合规场景评分失败: {str(e)}")

            # 选择得分最高的场景
            if any(scores.values()):
                selected_scenario = max(scores.items(), key=lambda x: x[1])[0]
                logger.info(f"场景检测结果: {scores}")
                logger.info(f"选择场景: {selected_scenario}")
                return selected_scenario

        # 如果场景适配模块不可用或评分全为0，使用基础规则检测
        _, ext = os.path.splitext(file_path)
        
        # 基于文件类型的基础规则
        if ext.lower() == '.csv':
            try:
                df = pd.read_csv(file_path)
                columns = [col.lower() for col in df.columns]
                if any(keyword in ' '.join(columns) for keyword in ['transaction', 'account', 'amount', 'balance', 'fraud', 'risk']):
                    return "fraud_detection"
                elif any(keyword in ' '.join(columns) for keyword in ['compliance', 'regulation', 'policy', 'rule']):
                    return "compliance"
            except:
                pass
        
        # 基于文件名的规则
        file_name = os.path.basename(file_path).lower()
        if any(kw in file_name for kw in ['customer', 'service', 'support', 'chat', 'qa']):
            return "customer_service"
        elif any(kw in file_name for kw in ['fraud', 'risk', 'transaction', 'alert']):
            return "fraud_detection"
        elif any(kw in file_name for kw in ['compliance', 'regulation', 'policy']):
            return "compliance"
        
        # 如果有内容，基于内容的规则
        if content:
            content_lower = content.lower()
            if any(kw in content_lower for kw in ['customer service', 'support', 'help', 'question', 'answer']):
                return "customer_service"
            elif any(kw in content_lower for kw in ['fraud', 'suspicious', 'risk', 'alert', 'transaction']):
                return "fraud_detection"
            elif any(kw in content_lower for kw in ['compliance', 'regulation', 'policy', 'requirement']):
                return "compliance"

        # 默认返回客服场景
        return "customer_service"

    except Exception as e:
        logger.error(f"场景检测失败: {str(e)}")
        return "customer_service"


def process_file(file_path: str, output_dir: str) -> Dict[str, Any]:
    """处理单个文件"""
    logger.info(f"\n{'='*50}\n处理文件: {file_path}\n{'='*50}")
    
    try:
        # 初始化处理器
        doc_processor = DocumentProcessor()
        chunk_manager = ChunkManager()
        info_processor = InformationProcessor()
        adaptive_system = EnhancedAdaptiveSystem()
        
        # 获取文件信息
        file_info = {
            'name': os.path.basename(file_path),
            'path': file_path,
            'size': os.path.getsize(file_path),
            'type': os.path.splitext(file_path)[1]
        }
        
        # 读取文档
        doc_content = doc_processor.process_document(file_path)
        if isinstance(doc_content, dict) and doc_content.get('total_pages'):
            file_info['total_pages'] = doc_content['total_pages']
            text_chunks = doc_content['text_chunks']
        else:
            text_chunks = chunk_manager.split_text(doc_content)
            file_info['total_pages'] = len(text_chunks)
        
        # 处理结果
        processed_chunks = []
        all_entities = []
        all_relations = []
        all_anomalies = []
        
        # 记录开始时间
        start_time = datetime.now()
        
        # 处理每个文本块
        for i, chunk in enumerate(text_chunks, 1):
            chunk_info = {
                **file_info,
                'current_page': i,
                'chunk_size': len(chunk)
            }
            
            # 使用自适应系统处理
            adaptive_result = adaptive_system.process(chunk, {'file_info': chunk_info})
            
            # 使用信息处理器处理
            info_result = info_processor.process(adaptive_result['text'], chunk_info)
            
            # 合并结果
            processed_chunk = {
                'text': chunk,
                'enhanced_text': adaptive_result['text'],
                'entities': info_result.get('entities', []),
                'relations': info_result.get('relations', []),
                'anomalies': info_result.get('anomalies', []),
                'scene': adaptive_result.get('scene'),
                'context_enhancements': adaptive_result.get('context_enhancements', []),
                'metadata': {
                    'chunk_number': i,
                    'chunk_size': len(chunk),
                    'processing_time': info_result.get('metadata', {}).get('processing_time', 0)
                }
            }
            
            processed_chunks.append(processed_chunk)
            all_entities.extend(processed_chunk['entities'])
            all_relations.extend(processed_chunk['relations'])
            all_anomalies.extend(processed_chunk['anomalies'])
            
            # 从处理结果中学习
            feedback = {
                'patterns': {
                    'company': [r'(?:[\u4e00-\u9fa5]+(?:股份|科技|信息|集团|控股))'],
                    'money': [r'(?:\d+(?:\.\d+)?(?:亿|万)?美金)']
                },
                'keywords': {
                    'financial': ['营收', '利润', '增长', '下滑'],
                    'tech': ['人工智能', '区块链', '云计算', '大数据']
                }
            }
            adaptive_system.learn_from_feedback(chunk, feedback)
        
        # 生成处理报告
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # 获取统计信息
        adaptive_stats = adaptive_system.get_statistics()
        info_stats = info_processor.get_statistics()
        
        # 保存结果
        result = {
            'file_info': file_info,
            'processing_summary': {
                'total_chunks': len(processed_chunks),
                'total_entities': len(all_entities),
                'total_relations': len(all_relations),
                'total_anomalies': len(all_anomalies),
                'processing_time': processing_time,
                'adaptive_system_stats': adaptive_stats,
                'information_processor_stats': info_stats
            },
            'processed_chunks': processed_chunks
        }
        
        # 保存到文件
        output_file = os.path.join(
            output_dir,
            f"{os.path.splitext(file_info['name'])[0]}_processed.json"
        )
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        logger.info(f"文件处理完成: {file_info['name']}")
        logger.info(f"处理时间: {processing_time:.2f} 秒")
        logger.info(f"发现实体: {len(all_entities)} 个")
        logger.info(f"发现关系: {len(all_relations)} 个")
        logger.info(f"发现异常: {len(all_anomalies)} 个")
        logger.info(f"结果已保存到: {output_file}")
        
        return result
        
    except Exception as e:
        logger.error(f"处理文件时出错: {str(e)}")
        logger.error(traceback.format_exc())
        return {
            'error': str(e),
            'file': file_path,
            'traceback': traceback.format_exc()
        }

def main(input_dir: str, output_dir: str):
    """主处理函数"""
    try:
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        # 获取所有文件
        files = []
        for root, _, filenames in os.walk(input_dir):
            for filename in filenames:
                if filename.endswith(('.pdf', '.docx', '.txt', '.csv')):
                    files.append(os.path.join(root, filename))
        
        logger.info(f"发现 {len(files)} 个文件待处理")
        
        # 处理所有文件
        results = []
        start_time = datetime.now()
        
        for file_path in files:
            result = process_file(file_path, output_dir)
            results.append(result)
        
        # 生成总体报告
        total_time = (datetime.now() - start_time).total_seconds()
        
        report = {
            'total_files': len(files),
            'processed_files': len([r for r in results if 'error' not in r]),
            'failed_files': len([r for r in results if 'error' in r]),
            'processing_time': total_time,
            'processed_file_list': [os.path.basename(f) for f in files],
            'failed_file_list': [
                os.path.basename(r['file']) 
                for r in results 
                if 'error' in r
            ],
            'summary': {
                'total_entities': sum(
                    r.get('processing_summary', {}).get('total_entities', 0)
                    for r in results
                    if 'error' not in r
                ),
                'total_relations': sum(
                    r.get('processing_summary', {}).get('total_relations', 0)
                    for r in results
                    if 'error' not in r
                ),
                'total_anomalies': sum(
                    r.get('processing_summary', {}).get('total_anomalies', 0)
                    for r in results
                    if 'error' not in r
                )
            }
        }
        
        # 保存报告
        report_file = os.path.join(output_dir, 'processing_report.json')
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        logger.info(f"\n{'='*50}")
        logger.info(f"处理完成")
        logger.info(f"总文件数: {report['total_files']}")
        logger.info(f"成功处理: {report['processed_files']}")
        logger.info(f"处理失败: {report['failed_files']}")
        logger.info(f"总处理时间: {total_time:.2f} 秒")
        logger.info(f"报告已保存到: {report_file}")
        
    except Exception as e:
        logger.error(f"处理过程出错: {str(e)}")
        logger.error(traceback.format_exc())
        sys.exit(1)

if __name__ == '__main__':
    # 设置默认目录
    default_input_dir = 'data'
    default_output_dir = 'output'
    
    if len(sys.argv) == 3:
        input_dir = sys.argv[1]
        output_dir = sys.argv[2]
    else:
        input_dir = default_input_dir
        output_dir = default_output_dir
        logger.info(f"使用默认目录 - 输入: {input_dir}, 输出: {output_dir}")
    
    if not os.path.exists(input_dir):
        logger.error(f"输入目录不存在: {input_dir}")
        sys.exit(1)
    
    main(input_dir, output_dir)